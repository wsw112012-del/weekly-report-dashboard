"""file_extract.py — 첨부 파일(docx/pdf/txt) 본문 텍스트 추출.

법률검토 문서를 입력받아 한국어 본문만 추출. 파일 자체는 서버 어디에도 저장하지 않고
메모리에서 처리. extract_text() 가 단일 진입점.
"""
from __future__ import annotations

import io
import re

_NBSP_RE = re.compile(r"[  ]")
_INVIS_RE = re.compile(r"[​‌‍⁠﻿]")


def _normalize_lines(text: str) -> str:
    """공통 정규화 — invisible/NBSP/다중 공백/빈 줄 묶음."""
    if not text:
        return ""
    text = _INVIS_RE.sub("", text)
    text = _NBSP_RE.sub(" ", text)
    # 줄 단위 공백 정리
    lines = []
    blank_pending = False
    for ln in text.split("\n"):
        ln = re.sub(r"[ \t]+", " ", ln).strip()
        if not ln:
            blank_pending = bool(lines)
            continue
        if blank_pending:
            lines.append("")
            blank_pending = False
        lines.append(ln)
    return "\n".join(lines)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # 표 안 텍스트도 포함
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t.strip())
    return "\n\n".join(parts)


def _extract_txt(data: bytes) -> str:
    # 인코딩 자동 감지 — utf-8 → cp949 fallback
    for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes, max_chars: int = 50000) -> str:
    """파일 확장자별 텍스트 추출 + 정규화 + max_chars cap.

    Raises:
        ValueError: 지원 안 하는 확장자 또는 추출 실패
    """
    name = (filename or "").lower()
    if name.endswith(".docx"):
        raw = _extract_docx(data)
    elif name.endswith(".pdf"):
        raw = _extract_pdf(data)
    elif name.endswith(".txt"):
        raw = _extract_txt(data)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {filename}")

    if not raw.strip():
        raise ValueError("파일에서 텍스트를 추출하지 못했습니다 (스캔 PDF 등 OCR 필요한 형식일 수 있음)")

    text = _normalize_lines(raw)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n... [길이 초과로 절단]"
    return text
